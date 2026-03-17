#!/usr/bin/env python3
"""
CV Tailor pipeline orchestrator — two-phase Python→Claude→Python handoff.

Phase 1 (prep):
  - Select base CV deterministically
  - Build analysis context → write data/pending-analysis.json
  - Stop. Claude reads pending-analysis.json and writes data/analysis.json.

Phase 2 (apply):
  - Read data/analysis.json (written by Claude)
  - Validate schema
  - Patch resume, generate cover letter, redline, QC, manifest
  - Clean up working files

Usage:
  python3 scripts/run_pipeline.py --phase prep --company Acme --role "AI PM" --jd-path /tmp/jd.txt
  python3 scripts/run_pipeline.py --phase apply --company Acme --role "AI PM"
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

# Direct imports — no subprocess
sys.path.insert(0, str(Path(__file__).resolve().parent))
from build_analysis import build as build_analysis_context, PENDING_PATH, ANALYSIS_PATH
from validate_analysis import validate as validate_analysis
from docx_safe_patch import apply_safe_patch
from generate_redline import generate as generate_redline
from quality_gate import qc as run_quality_gate
from index_store import rebuild_registry, CV_BASE

SCRIPT_DIR = Path(__file__).resolve().parent


def _role_clean(role: str) -> str:
    return re.sub(r'[^\w\s-]', '', role).strip().replace(' ', '_')[:60]


def _next_version(company_dir: Path, rc: str, date_str: str) -> int:
    nums = []
    # Scan all versions for this role regardless of date, so revisions
    # always increment even across date boundaries.
    for p in company_dir.glob(f'Manifest_{rc}_*_v*.json'):
        m = re.search(r'_v(\d+)\.json$', p.name)
        if m:
            nums.append(int(m.group(1)))
    return (max(nums) + 1) if nums else 1


def _read_doc_text(path: Path) -> str:
    d = Document(str(path))
    return '\n'.join([p.text for p in d.paragraphs if p.text])


def _build_cover_letter(out_path: Path, company: str, role: str, paragraphs: list[str]) -> None:
    d = Document()
    normal = d.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    d.add_paragraph(datetime.now().strftime('%B %d, %Y'))
    d.add_paragraph('')
    d.add_paragraph(f'{company} Hiring Team')
    d.add_paragraph(f'Re: {role}')
    d.add_paragraph('')
    d.add_paragraph('Dear Hiring Team,')
    d.add_paragraph('')
    for p in paragraphs:
        d.add_paragraph(p)
        d.add_paragraph('')
    d.add_paragraph('Best regards,')
    d.add_paragraph('[YOUR_FULL_NAME]')

    for para in d.paragraphs:
        for run in para.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
    d.save(out_path)


def _write_change_receipt(path: Path, expected_edits: list[dict], patch_result: dict) -> None:
    lines = [
        'CV Tailor Change Receipt',
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Expected edits: {len(expected_edits)}",
        f"Applied edits: {patch_result.get('changed_count', 0)}",
        '',
        'Expected edits:',
    ]
    for i, e in enumerate(expected_edits, 1):
        lines.append(f"{i}. OLD: {e.get('old', '')}")
        lines.append(f"   NEW: {e.get('new', '')}")
    lines += ['', 'Applied edit events:']
    for c in patch_result.get('changes', []):
        lines.append(f"- paragraph {c.get('paragraph')} | OLD: {c.get('old')} | NEW: {c.get('new')}")
    path.write_text('\n'.join(lines), encoding='utf-8')


def cmd_prep(company: str, role: str, jd_path: str) -> int:
    """Phase 1: select base CV, build analysis context, write pending-analysis.json."""
    jd = Path(jd_path)
    if not jd.exists():
        print(f'ERROR: JD file missing: {jd}')
        return 1

    rebuild_registry(CV_BASE)
    # Import here to avoid circular at module level
    from select_base_cv import select as select_base_cv
    choice = select_base_cv(role, company)
    base_path = choice.get('selected_base_cv') or choice.get('master_cv_candidate') or ''
    base = Path(base_path)
    if not base.exists():
        master_dir = CV_BASE / 'Master CV'
        candidates = sorted([p for p in master_dir.glob('*.docx') if 'cover' not in p.name.lower()])
        if candidates:
            base = candidates[0]
        else:
            print(f'ERROR: base resume not found: {base_path}')
            return 1

    jd_text = jd.read_text(encoding='utf-8')
    build_analysis_context(company, role, str(base), jd_text)

    print(f'\n[prep] Done. Next steps:')
    print(f'  1. Read {PENDING_PATH}')
    print(f'  2. Evaluate and write edits to {ANALYSIS_PATH}')
    print(f'  3. Run: python3 scripts/run_pipeline.py --phase apply --company "{company}" --role "{role}"')
    return 0


def cmd_apply(company: str, role: str) -> dict:
    """Phase 2: read analysis.json, apply patches, generate all artifacts."""
    if not ANALYSIS_PATH.exists():
        print(f'ERROR: {ANALYSIS_PATH} not found. Run --phase prep and have Claude write analysis.json.')
        sys.exit(1)

    analysis = json.loads(ANALYSIS_PATH.read_text(encoding='utf-8'))

    errs = validate_analysis(analysis)
    if errs:
        print('ERROR: analysis.json failed validation:')
        for e in errs:
            print(f'  - {e}')
        sys.exit(1)

    base = Path(analysis.get('base_resume_path', ''))
    if not base.exists():
        print(f'ERROR: base resume from analysis.json missing: {base}')
        sys.exit(1)

    company_dir = CV_BASE / company
    company_dir.mkdir(parents=True, exist_ok=True)
    date_str = datetime.now().strftime('%Y%m%d')
    rc = _role_clean(role)
    ver = _next_version(company_dir, rc, date_str)
    vs = f'v{ver:03d}'

    resume_path = company_dir / f'Resume_Tailored_{rc}_{date_str}_{vs}.docx'
    cover_path = company_dir / f'Cover_Letter_{rc}_{date_str}_{vs}.docx'
    redline_path = company_dir / f'Redline_{rc}_{date_str}_{vs}.docx'
    qc_path = company_dir / f'QC_{rc}_{date_str}_{vs}.json'
    manifest_path = company_dir / f'Manifest_{rc}_{date_str}_{vs}.json'
    receipt_path = company_dir / f'Changes_Applied_{rc}_{date_str}_{vs}.txt'

    edits = analysis.get('summary_edits', []) + analysis.get('bullet_edits', [])

    patch_result = apply_safe_patch(base, resume_path, edits)
    _write_change_receipt(receipt_path, edits, patch_result)
    _build_cover_letter(cover_path, company, role, analysis.get('cover_letter_paragraphs', []))

    # generate_redline reads analysis from file — pass the analysis path
    generate_redline(str(base), str(ANALYSIS_PATH), str(redline_path))

    qc_result = run_quality_gate(str(resume_path), str(redline_path))
    qc_path.write_text(json.dumps(qc_result, indent=2))

    resume_text = _read_doc_text(resume_path)
    verified = []
    for e in edits:
        old = e.get('old', '')
        new = e.get('new', '')
        verified.append({
            'old': old,
            'new': new,
            'new_present': bool(new and new in resume_text),
            'old_still_present': bool(old and old in resume_text),
        })

    verified_count = sum(1 for v in verified if v['new_present'])
    expected_count = len(edits)
    min_required = min(5, expected_count) if expected_count > 0 else 0

    fail_reasons = []
    if qc_result.get('status') != 'pass':
        fail_reasons.append('qc_failed')
    if patch_result.get('changed_count', 0) < min_required:
        fail_reasons.append(f"edit_count_below_threshold:{patch_result.get('changed_count', 0)}<{min_required}")
    if verified_count < min_required:
        fail_reasons.append(f'verified_edits_below_threshold:{verified_count}<{min_required}')

    manifest = {
        'status': 'pass' if not fail_reasons else 'fail',
        'company': company,
        'role': role,
        'base_resume_used': str(base),
        'version': vs,
        'expected_edits': expected_count,
        'patch_changes': patch_result.get('changed_count', 0),
        'verified_edits': verified_count,
        'files': {
            'resume': str(resume_path),
            'cover_letter': str(cover_path),
            'redline': str(redline_path),
            'qc': str(qc_path),
            'changes_receipt': str(receipt_path),
        },
        'qc': qc_result,
        'verification': verified,
        'fail_reasons': fail_reasons,
    }
    manifest_path.write_text(json.dumps(manifest, indent=2))

    # Clean up working files
    for p in (ANALYSIS_PATH, PENDING_PATH):
        if p.exists():
            p.unlink()

    rebuild_registry(CV_BASE)
    return manifest


def main() -> int:
    ap = argparse.ArgumentParser(description='CV Tailor pipeline — two-phase')
    ap.add_argument('--phase', choices=['prep', 'apply'], required=True)
    ap.add_argument('--company', required=True)
    ap.add_argument('--role', required=True)
    ap.add_argument('--jd-path', help='Path to JD text file — required for --phase prep')
    args = ap.parse_args()

    if args.phase == 'prep':
        if not args.jd_path:
            print('ERROR: --jd-path required for --phase prep')
            return 1
        return cmd_prep(args.company, args.role, args.jd_path)

    if args.phase == 'apply':
        m = cmd_apply(args.company, args.role)
        print(json.dumps(m, indent=2))
        return 0 if m['status'] == 'pass' else 1

    return 0


if __name__ == '__main__':
    raise SystemExit(main())
