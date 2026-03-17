#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path
from datetime import datetime
from docx import Document


def generate(base_path: str, analysis_path: str, out_path: str) -> None:
    analysis = json.loads(Path(analysis_path).read_text())
    d = Document()
    d.add_heading('Resume Redline', level=1)
    d.add_paragraph(f'Base: {base_path}')
    d.add_paragraph(f'Analysis: {analysis_path}')
    d.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')

    changes = analysis.get('summary_edits', []) + analysis.get('bullet_edits', [])
    d.add_paragraph(f'Total planned edits: {len(changes)}')

    for i, ch in enumerate(changes, 1):
        d.add_heading(f'Change {i}', level=2)
        d.add_paragraph('Before:')
        d.add_paragraph(ch.get('old', ''))
        d.add_paragraph('After:')
        d.add_paragraph(ch.get('new', ''))

    d.save(out_path)


if __name__ == '__main__':
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument('--base', required=True)
    ap.add_argument('--analysis', required=True)
    ap.add_argument('--out', required=True)
    args = ap.parse_args()
    generate(args.base, args.analysis, args.out)
    print(args.out)
