#!/usr/bin/env python3
from __future__ import annotations

import json
from pathlib import Path
from docx import Document

"""
Run-aware safe patcher.
- Finds replacements at paragraph text level (across runs)
- Re-distributes updated text back into existing runs to preserve run styling layout as much as possible
"""


def _redistribute_text_to_runs(paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(new_text)
        return

    original_lengths = [len(r.text or "") for r in runs]
    cursor = 0
    for i, r in enumerate(runs):
        take = original_lengths[i]
        # last run gets remainder
        if i == len(runs) - 1:
            r.text = new_text[cursor:]
        else:
            r.text = new_text[cursor:cursor + take]
            cursor += take


def apply_safe_patch(src: Path, dst: Path, repls: list[dict]) -> dict:
    d = Document(str(src))
    changed = []

    for para_idx, p in enumerate(d.paragraphs):
        full = "".join(r.text or "" for r in p.runs) if p.runs else p.text or ""
        if not full:
            continue

        updated = full
        local_changes = []
        for rep in repls:
            old = rep.get('old', '')
            new = rep.get('new', '')
            if old and old in updated:
                updated = updated.replace(old, new, 1)
                local_changes.append({'old': old, 'new': new})

        if updated != full:
            _redistribute_text_to_runs(p, updated)
            for ch in local_changes:
                changed.append({
                    'paragraph': para_idx,
                    'old': ch['old'],
                    'new': ch['new']
                })

    d.save(str(dst))
    return {'changed_count': len(changed), 'changes': changed}


if __name__ == '__main__':
    import sys
    if len(sys.argv) != 4:
        print('Usage: docx_safe_patch.py <input.docx> <output.docx> <replacements.json>')
        raise SystemExit(1)
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    cfg = json.loads(Path(sys.argv[3]).read_text())
    result = apply_safe_patch(src, dst, cfg.get('replacements', []))
    print(json.dumps(result, indent=2))
